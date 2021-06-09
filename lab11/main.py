import math
import re
import shutil
import matplotlib.pyplot as plt
import PIL
from PIL import Image
import docx
import requests
from docx.shared import RGBColor

title_parse = re.compile(r'\b(Title)\b:(.*)')
author_parse = re.compile(r'\b(Author)\b:(.*)')


def run():

    results = file_reader('11-0.txt')
    # print(results)
    my_photo = download_photo()
    new_photo = manipulate_photo(my_photo, 'logo.png')

    createDoc(str(results[1]), str(results[0]), new_photo, results[2])


def download_photo():
    image_url = "https://cdn-ssl.s7.disneystore.com/is/image/DisneyShopping/7745055550016?fmt=webp&qlt=90&wid=652&hei=652"
    filename = 'myphoto.jpg'
    r = requests.get(image_url, stream=True)
    if r.status_code == 200:
        r.raw.decode_content = True
        with open(filename, 'wb') as f:
            shutil.copyfileobj(r.raw, f)

        print('Image sucessfully Downloaded: ', filename)
    else:
        print('Image Couldn\'t be retreived')

    return filename


def manipulate_photo(filename, logoname):
    img = Image.open(filename)
    img2 = img.crop((0, 0, 550, 400))
    # coordinates, (left,top,right,bottom), left<right, top<bottom
    new1_img = img2.resize((300, 200), PIL.Image.ANTIALIAS)
    logo = Image.open(logoname)
    logo_modified = logo.resize((50, 50), PIL.Image.ANTIALIAS).rotate(45)
    new1_img.paste(logo_modified, (200, 0))
    name = 'modified.jpg'
    new1_img.save(name)
    return name


def file_reader(new_file):
    num = 0
    chapterList = []
    word_counter = 0
    paraghraphCounter = 1
    deleteTitle = True
    paragraph_length_list = []

    try:
        with open(new_file, "r", encoding='UTF8') as file:
            for line in file:
                parser_tit = re.match(title_parse, line)
                if parser_tit is not None:
                    title = str(parser_tit.group(2))
                    print(title)
                parser_aut = re.match(author_parse, line)
                if parser_aut is not None:
                    author = parser_aut.group(2)
                    print(author)
                if num == 3:
                    if not deleteTitle:
                        if 'CHAPTER II.' not in line:
                            chapterList.append(line)
                            if line != "\n" and '*' not in line:
                                word_counter = word_counter + len(line.split())
                            else:
                                if word_counter != 0:

                                    word_counter = word_counter - word_counter % 10
                                    paragraph_length_list.append(word_counter)
                                    print(f'Paragraph {paraghraphCounter} and rounded num of words is {word_counter}')
                                    word_counter = 0
                                    paraghraphCounter = paraghraphCounter + 1
                    else:
                        deleteTitle = False
                if 'CHAPTER I.' in line:
                    num = num+1
                if 'CHAPTER II.' in line:
                    num = num+1

            # for line in chapterList:
            #     print(line)
            littlecounter = 1
            counting_dictionary = {}
            for number in sorted(paragraph_length_list):
                for num2 in paragraph_length_list:
                    if number == num2:
                        littlecounter += 1
                counting_dictionary[number] = littlecounter
                littlecounter = 0
            print(f'How many times a value is in there:{ counting_dictionary}')
            print(f'sorted list: {sorted(paragraph_length_list)}')

    except FileNotFoundError:
        print('File not found')
        exit()
    return author, title, paragraph_length_list


def createDoc(title, author, picture, params):
    doc = docx.Document()
    doc.add_heading(title, 0)
    doc.add_heading(f'By: {author}', 1)
    p = doc.add_paragraph()
    runner = p.add_run('prepared by Julia Moska')
    runner.italic = True
    doc.add_picture(r'C:\Users\julia\Documents\GitHub\python\lab11\\' + str(picture))
    doc.add_page_break()
    doc.add_heading('Plot page:', 0)
    plt.bar(range(1, len(params)+1), params)
    plt.xlabel('paragraph no.')
    plt.ylabel('Number of words')
    plt.title('Words per paragraph')
    name = 'firstplot.png'
    plt.savefig('firstplot.png')
    doc.add_picture(r'C:\Users\julia\Documents\GitHub\python\lab11\\'+name)

    p = doc.add_paragraph()
    runner = p.add_run(f'Total Number of paragraphs: {len(params)} ')
    runner.bold = True
    para = doc.add_paragraph().add_run(f'Maximum number of words: {max(params)} ')
    para.font.color.rgb = RGBColor(124, 252, 0)
    param = doc.add_paragraph().add_run(f'Minimum number of words: {min(params)} ')
    param.font.color.rgb = RGBColor(255, 0, 0)
    param.italic = True
    sum = 0
    for elem in params:
        sum = sum+int(elem)
    doc.add_paragraph(f'Average number of words: {math.floor(sum/len(params))} ')
    doc.save('mydocument.docx')


if __name__ == '__main__':
    run()
